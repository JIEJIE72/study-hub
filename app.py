from flask import Flask, render_template, request, redirect, url_for, send_from_directory, send_file, session
from docx import Document
from datetime import datetime
import os
import sqlite3
import io
from bs4 import BeautifulSoup

app = Flask(__name__)

# ---------------- LOGIN SYSTEM ----------------
app.secret_key = "studyhub_secret_2026"

USERNAME = "aaron"
PASSWORD = "1234"

def login_required():
    return session.get("logged_in")

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None

    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        if username == USERNAME and password == PASSWORD:
            session["logged_in"] = True
            return redirect("/")
        else:
            error = "Invalid login"

    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    session.pop("logged_in", None)
    return redirect("/login")

# ---------------- PATHS ----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "notes.db")
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ---------------- DB INIT ----------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("""
        CREATE TABLE IF NOT EXISTS topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE,
            order_index INTEGER DEFAULT 0
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic_id INTEGER,
            title TEXT,
            subject TEXT,
            content TEXT,
            order_index INTEGER DEFAULT 0
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            note_id INTEGER,
            filename TEXT
        )
    """)

    conn.commit()
    conn.close()

init_db()

# ---------------- HOME ----------------
@app.route("/")
def index():
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT * FROM topics ORDER BY order_index ASC")
    topics = c.fetchall()
    conn.close()
    return render_template("index.html", topics=topics)

# ---------------- TOPIC ----------------
@app.route("/topic/<int:id>")
def topic(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT name FROM topics WHERE id=?", (id,))
    topic = c.fetchone()

    if not topic:
        conn.close()
        return "Topic not found", 404

    c.execute("""
        SELECT * FROM notes
        WHERE topic_id=?
        ORDER BY order_index ASC
    """, (id,))
    notes = c.fetchall()

    conn.close()

    return render_template("topic.html", topic=topic, notes=notes, topic_id=id)

# ---------------- ADD TOPIC ----------------
@app.route("/add_topic", methods=["POST"])
def add_topic():
    if not login_required():
        return redirect("/login")

    name = request.form.get("name", "").strip()
    if not name:
        return redirect("/")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT MAX(order_index) FROM topics")
    max_order = c.fetchone()[0]
    next_order = (max_order + 1) if max_order is not None else 1

    c.execute("""
        INSERT INTO topics (name, order_index)
        VALUES (?, ?)
    """, (name, next_order))

    conn.commit()
    conn.close()
    return redirect("/")

# ---------------- ADD NOTE ----------------
@app.route("/add/<int:topic_id>", methods=["GET", "POST"])
def add(topic_id):
    if not login_required():
        return redirect("/login")

    if request.method == "POST":

        title = request.form.get("title", "")
        subject = request.form.get("subject", "")
        content = request.form.get("content", "")

        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()

        c.execute("SELECT MAX(order_index) FROM notes WHERE topic_id=?", (topic_id,))
        max_order = c.fetchone()[0]
        next_order = (max_order + 1) if max_order is not None else 1

        c.execute("""
            INSERT INTO notes (topic_id, title, subject, content, order_index)
            VALUES (?, ?, ?, ?, ?)
        """, (topic_id, title, subject, content, next_order))

        note_id = c.lastrowid

        files = request.files.getlist("file")
        for f in files:
            if f and f.filename:
                filename = os.path.basename(f.filename)
                path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                f.save(path)

                c.execute("INSERT INTO files (note_id, filename) VALUES (?, ?)",
                          (note_id, filename))

        conn.commit()
        conn.close()

        return redirect(url_for("topic", id=topic_id))

    return render_template("add.html", topic_id=topic_id)

# ---------------- VIEW NOTE ----------------
@app.route("/note/<int:id>")
def view(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT * FROM notes WHERE id=?", (id,))
    note = c.fetchone()

    c.execute("SELECT * FROM files WHERE note_id=?", (id,))
    files = c.fetchall()

    conn.close()

    return render_template("view.html", note=note, files=files, topic_id=note[1])

# ---------------- EDIT NOTE ----------------
@app.route("/edit/<int:id>", methods=["GET", "POST"])
def edit(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    if request.method == "POST":

        title = request.form.get("title", "")
        subject = request.form.get("subject", "")
        content = request.form.get("content", "")

        c.execute("""
            UPDATE notes
            SET title=?, subject=?, content=?
            WHERE id=?
        """, (title, subject, content, id))

        files = request.files.getlist("file")
        for f in files:
            if f and f.filename:
                filename = os.path.basename(f.filename)
                path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                f.save(path)

                c.execute("INSERT INTO files (note_id, filename) VALUES (?, ?)",
                          (id, filename))

        conn.commit()
        conn.close()

        return redirect(url_for("view", id=id))

    c.execute("SELECT * FROM notes WHERE id=?", (id,))
    note = c.fetchone()

    conn.close()

    return render_template("edit.html", note=note)

# ---------------- MOVE TOPIC ----------------
@app.route("/topic_up/<int:id>")
def topic_up(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT id, order_index FROM topics WHERE id=?", (id,))
    current = c.fetchone()

    c.execute("""
        SELECT id, order_index FROM topics
        WHERE order_index < ?
        ORDER BY order_index DESC LIMIT 1
    """, (current[1],))

    above = c.fetchone()

    if above:
        c.execute("UPDATE topics SET order_index=? WHERE id=?",
                  (above[1], current[0]))
        c.execute("UPDATE topics SET order_index=? WHERE id=?",
                  (current[1], above[0]))

    conn.commit()
    conn.close()
    return redirect("/")

@app.route("/topic_down/<int:id>")
def topic_down(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT id, order_index FROM topics WHERE id=?", (id,))
    current = c.fetchone()

    c.execute("""
        SELECT id, order_index FROM topics
        WHERE order_index > ?
        ORDER BY order_index ASC LIMIT 1
    """, (current[1],))

    below = c.fetchone()

    if below:
        c.execute("UPDATE topics SET order_index=? WHERE id=?",
                  (below[1], current[0]))
        c.execute("UPDATE topics SET order_index=? WHERE id=?",
                  (current[1], below[0]))

    conn.commit()
    conn.close()
    return redirect("/")

# ---------------- MOVE NOTE ----------------
@app.route("/move_up/<int:id>/<int:topic_id>")
def move_up(id, topic_id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT id, order_index FROM notes WHERE id=?", (id,))
    current = c.fetchone()

    c.execute("""
        SELECT id, order_index FROM notes
        WHERE topic_id=? AND order_index < ?
        ORDER BY order_index DESC LIMIT 1
    """, (topic_id, current[1]))

    above = c.fetchone()

    if above:
        c.execute("UPDATE notes SET order_index=? WHERE id=?",
                  (above[1], current[0]))
        c.execute("UPDATE notes SET order_index=? WHERE id=?",
                  (current[1], above[0]))

    conn.commit()
    conn.close()
    return redirect(url_for("topic", id=topic_id))

@app.route("/move_down/<int:id>/<int:topic_id>")
def move_down(id, topic_id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT id, order_index FROM notes WHERE id=?", (id,))
    current = c.fetchone()

    c.execute("""
        SELECT id, order_index FROM notes
        WHERE topic_id=? AND order_index > ?
        ORDER BY order_index ASC LIMIT 1
    """, (topic_id, current[1]))

    below = c.fetchone()

    if below:
        c.execute("UPDATE notes SET order_index=? WHERE id=?",
                  (below[1], current[0]))
        c.execute("UPDATE notes SET order_index=? WHERE id=?",
                  (current[1], below[0]))

    conn.commit()
    conn.close()
    return redirect(url_for("topic", id=topic_id))

# ---------------- RENAME TOPIC ----------------
@app.route("/rename_topic/<int:id>", methods=["POST"])
def rename_topic(id):
    if not login_required():
        return redirect("/login")

    name = request.form.get("name", "").strip()

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE topics SET name=? WHERE id=?", (name, id))
    conn.commit()
    conn.close()

    return redirect(url_for("topic", id=id))

# ---------------- DELETE TOPIC ----------------
@app.route("/delete_topic/<int:id>", methods=["POST"])
def delete_topic(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT COUNT(*) FROM notes WHERE topic_id=?", (id,))
    count = c.fetchone()[0]
    conn.close()

    if count > 0:
        return redirect(url_for("topic", id=id))

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM topics WHERE id=?", (id,))
    conn.commit()
    conn.close()

    return redirect("/")

# ---------------- DELETE NOTE ----------------
@app.route("/delete/<int:id>", methods=["POST"])
def delete(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT topic_id FROM notes WHERE id=?", (id,))
    topic_id = c.fetchone()[0]

    c.execute("DELETE FROM files WHERE note_id=?", (id,))
    c.execute("DELETE FROM notes WHERE id=?", (id,))

    conn.commit()
    conn.close()

    return redirect(url_for("topic", id=topic_id))

# ---------------- DELETE FILE ----------------
@app.route("/delete_file/<int:file_id>/<int:note_id>", methods=["POST"])
def delete_file(file_id, note_id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT filename FROM files WHERE id=?", (file_id,))
    row = c.fetchone()

    if row:
        path = os.path.join(UPLOAD_FOLDER, row[0])
        if os.path.exists(path):
            os.remove(path)

        c.execute("DELETE FROM files WHERE id=?", (file_id,))

    conn.commit()
    conn.close()

    return redirect(url_for("view", id=note_id))

# ---------------- FILE SERVE ----------------
@app.route("/uploads/<filename>")
def uploads(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)

# ---------------- EXPORT WORD ----------------
@app.route("/export/<int:id>")
def export(id):
    if not login_required():
        return redirect("/login")

    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT title, subject, content FROM notes WHERE id=?", (id,))
    note = c.fetchone()
    conn.close()

    title, subject, content = note

    doc = Document()
    doc.add_heading(title or "", 0)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Subject:\n")
    r1.bold = True
    p1.add_run(subject or "")

    p2 = doc.add_paragraph()
    r2 = p2.add_run("\nContents:\n")
    r2.bold = True

    soup = BeautifulSoup(content or "", "html.parser")
    text = soup.get_text("\n")
    p2.add_run(text)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{title}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ---------------- RUN ----------------
if __name__ == "__main__":
    app.run()